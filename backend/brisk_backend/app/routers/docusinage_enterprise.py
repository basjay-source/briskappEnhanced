from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, BackgroundTasks, Query
from fastapi.responses import FileResponse, JSONResponse
from typing import List, Optional, Dict, Any, Union
import os
import uuid
import asyncio
import json
from datetime import datetime
import tempfile
import shutil
import mimetypes
from pathlib import Path

try:
    import cv2
    import numpy as np
    from PIL import Image, ImageEnhance, ImageFilter
    import pytesseract
    from pdf2image import convert_from_path
    import fitz  # PyMuPDF
    import pandas as pd
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from pptx import Presentation
    from pptx.util import Inches as PptxInches
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    import camelot
    import tabula
    OCR_AVAILABLE = True
except ImportError as e:
    print(f"OCR dependencies not available: {e}")
    OCR_AVAILABLE = False

from app.middleware.auth import get_current_user

router = APIRouter()

UPLOAD_DIR = "/tmp/docusinage_enterprise_uploads"
PROCESSED_DIR = "/tmp/docusinage_enterprise_processed"
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(PROCESSED_DIR, exist_ok=True)

SUPPORTED_FORMATS = {
    'input': ['.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif', '.webp', 
              '.xlsx', '.xls', '.csv', '.docx', '.doc', '.pptx', '.ppt', '.txt', 
              '.html', '.json', '.xml', '.rtf'],
    'output': ['.pdf', '.xlsx', '.csv', '.docx', '.pptx', '.txt', '.html', '.json', 
               '.png', '.jpg', '.xml', '.rtf']
}

class EnterpriseDocumentProcessor:
    """Enterprise-grade document processor with comprehensive OCR and conversion capabilities"""
    
    def __init__(self):
        self.temp_files = []
        self.processing_status = {}
    
    def cleanup_temp_files(self):
        """Clean up temporary files"""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception:
                pass
        self.temp_files.clear()
    
    def enhance_image_for_ocr(self, image_path: str) -> str:
        """Advanced image enhancement for better OCR results"""
        if not OCR_AVAILABLE:
            return image_path
            
        try:
            img = cv2.imread(image_path)
            if img is None:
                return image_path
            
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            denoised = cv2.fastNlMeansDenoising(gray)
            
            thresh = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            
            kernel = np.ones((1, 1), np.uint8)
            cleaned = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
            
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(cleaned)
            
            enhanced_path = image_path.replace('.', '_enhanced.')
            cv2.imwrite(enhanced_path, enhanced)
            self.temp_files.append(enhanced_path)
            
            return enhanced_path
            
        except Exception as e:
            print(f"Image enhancement failed: {e}")
            return image_path
    
    def extract_text_with_ocr(self, image_path: str) -> Dict[str, Any]:
        """Extract text from image using advanced OCR with confidence scoring"""
        if not OCR_AVAILABLE:
            return {"text": "", "confidence": 0, "words": [], "lines": []}
        
        try:
            enhanced_path = self.enhance_image_for_ocr(image_path)
            
            configs = [
                r'--oem 3 --psm 6',  # Uniform block of text
                r'--oem 3 --psm 4',  # Single column of text
                r'--oem 3 --psm 3',  # Fully automatic page segmentation
                r'--oem 3 --psm 1',  # Automatic page segmentation with OSD
            ]
            
            best_result = {"text": "", "confidence": 0, "words": [], "lines": []}
            
            for config in configs:
                try:
                    data = pytesseract.image_to_data(enhanced_path, config=config, output_type=pytesseract.Output.DICT)
                    
                    text_blocks = []
                    words_data = []
                    lines_data = []
                    total_confidence = 0
                    valid_words = 0
                    
                    current_line = {"text": "", "bbox": None, "words": []}
                    last_line_num = -1
                    
                    for i in range(len(data['text'])):
                        if int(data['conf'][i]) > 30:  # Filter low confidence words
                            word_info = {
                                'text': data['text'][i],
                                'confidence': int(data['conf'][i]),
                                'bbox': {
                                    'x': data['left'][i],
                                    'y': data['top'][i],
                                    'width': data['width'][i],
                                    'height': data['height'][i]
                                }
                            }
                            words_data.append(word_info)
                            text_blocks.append(data['text'][i])
                            total_confidence += int(data['conf'][i])
                            valid_words += 1
                            
                            line_num = data['line_num'][i]
                            if line_num != last_line_num:
                                if current_line["text"]:
                                    lines_data.append(current_line)
                                current_line = {"text": data['text'][i], "bbox": word_info['bbox'], "words": [word_info]}
                                last_line_num = line_num
                            else:
                                current_line["text"] += " " + data['text'][i]
                                current_line["words"].append(word_info)
                    
                    if current_line["text"]:
                        lines_data.append(current_line)
                    
                    full_text = ' '.join(text_blocks)
                    avg_confidence = total_confidence / valid_words if valid_words > 0 else 0
                    
                    result = {
                        "text": full_text,
                        "confidence": avg_confidence,
                        "words": words_data,
                        "lines": lines_data,
                        "total_words": valid_words
                    }
                    
                    if avg_confidence > best_result["confidence"]:
                        best_result = result
                        
                except Exception as e:
                    continue
            
            return best_result
            
        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return {"text": "", "confidence": 0, "words": [], "lines": []}
    
    def extract_tables_from_pdf(self, pdf_path: str) -> List[Dict]:
        """Extract tables from PDF using multiple advanced methods"""
        tables = []
        
        if not OCR_AVAILABLE:
            return tables
        
        try:
            try:
                camelot_tables = camelot.read_pdf(pdf_path, pages='all', flavor='lattice')
                for i, table in enumerate(camelot_tables):
                    if table.accuracy > 50:
                        tables.append({
                            'method': 'camelot_lattice',
                            'page': table.page,
                            'accuracy': table.accuracy,
                            'data': table.df.to_dict('records'),
                            'headers': table.df.columns.tolist(),
                            'shape': table.df.shape
                        })
            except Exception as e:
                print(f"Camelot lattice extraction failed: {e}")
            
            try:
                camelot_stream = camelot.read_pdf(pdf_path, pages='all', flavor='stream')
                for i, table in enumerate(camelot_stream):
                    if table.accuracy > 40:
                        tables.append({
                            'method': 'camelot_stream',
                            'page': table.page,
                            'accuracy': table.accuracy,
                            'data': table.df.to_dict('records'),
                            'headers': table.df.columns.tolist(),
                            'shape': table.df.shape
                        })
            except Exception as e:
                print(f"Camelot stream extraction failed: {e}")
            
            try:
                tabula_tables = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True)
                for i, df in enumerate(tabula_tables):
                    if not df.empty and len(df.columns) > 1:
                        tables.append({
                            'method': 'tabula',
                            'page': i + 1,
                            'accuracy': 75,
                            'data': df.to_dict('records'),
                            'headers': df.columns.tolist(),
                            'shape': df.shape
                        })
            except Exception as e:
                print(f"Tabula extraction failed: {e}")
            
            if not tables:
                try:
                    images = convert_from_path(pdf_path, dpi=300)
                    for page_num, image in enumerate(images, 1):
                        temp_image_path = f"/tmp/pdf_page_{page_num}_{uuid.uuid4().hex}.png"
                        image.save(temp_image_path)
                        self.temp_files.append(temp_image_path)
                        
                        ocr_result = self.extract_text_with_ocr(temp_image_path)
                        
                        lines = ocr_result.get('lines', [])
                        table_data = []
                        
                        for line in lines:
                            text = line.get('text', '').strip()
                            if text:
                                row = [cell.strip() for cell in text.split('  ') if cell.strip()]
                                if len(row) > 1:
                                    table_data.append(row)
                        
                        if table_data and len(table_data) > 1:
                            max_cols = max(len(row) for row in table_data)
                            normalized_data = []
                            for row in table_data:
                                normalized_row = row + [''] * (max_cols - len(row))
                                normalized_data.append(normalized_row)
                            
                            headers = [f"Column_{i+1}" for i in range(max_cols)]
                            if normalized_data:
                                headers = normalized_data[0] if len(normalized_data[0]) == max_cols else headers
                                data_rows = normalized_data[1:] if len(normalized_data) > 1 else normalized_data
                                
                                df_data = []
                                for row in data_rows:
                                    row_dict = {headers[i]: row[i] if i < len(row) else '' for i in range(len(headers))}
                                    df_data.append(row_dict)
                                
                                tables.append({
                                    'method': 'ocr_table_detection',
                                    'page': page_num,
                                    'accuracy': ocr_result.get('confidence', 0),
                                    'data': df_data,
                                    'headers': headers,
                                    'shape': (len(df_data), len(headers))
                                })
                except Exception as e:
                    print(f"OCR table extraction failed: {e}")
            
        except Exception as e:
            print(f"Table extraction failed: {e}")
        
        return tables
    
    def convert_pdf_to_excel(self, pdf_path: str, output_path: str) -> Dict[str, Any]:
        """Convert PDF to Excel with advanced table structure preservation"""
        try:
            tables = self.extract_tables_from_pdf(pdf_path)
            
            if not tables:
                return self.convert_pdf_to_excel_ocr_fallback(pdf_path, output_path)
            
            wb = Workbook()
            wb.remove(wb.active)
            
            total_accuracy = 0
            sheets_created = 0
            
            for i, table_info in enumerate(tables):
                sheet_name = f"Table_{i+1}_Page_{table_info['page']}"[:31]  # Excel sheet name limit
                ws = wb.create_sheet(title=sheet_name)
                
                headers = table_info['headers']
                data = table_info['data']
                
                header_fill = PatternFill(start_color="FF6B35", end_color="FF6B35", fill_type="solid")
                header_font = Font(bold=True, color="FFFFFF")
                
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=str(header))
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center')
                
                for row_idx, row_data in enumerate(data, 2):
                    for col_idx, header in enumerate(headers, 1):
                        value = row_data.get(header, '')
                        cell = ws.cell(row=row_idx, column=col_idx, value=str(value))
                        cell.alignment = Alignment(wrap_text=True)
                
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column_letter].width = adjusted_width
                
                total_accuracy += table_info.get('accuracy', 0)
                sheets_created += 1
            
            wb.save(output_path)
            
            avg_accuracy = total_accuracy / sheets_created if sheets_created > 0 else 0
            
            return {
                'success': True,
                'tables_found': len(tables),
                'sheets_created': sheets_created,
                'quality_score': avg_accuracy,
                'method': 'advanced_table_extraction'
            }
            
        except Exception as e:
            print(f"PDF to Excel conversion failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def convert_pdf_to_excel_ocr_fallback(self, pdf_path: str, output_path: str) -> Dict[str, Any]:
        """Fallback OCR conversion when table extraction fails"""
        try:
            images = convert_from_path(pdf_path, dpi=300)
            
            wb = Workbook()
            wb.remove(wb.active)
            
            total_confidence = 0
            pages_processed = 0
            
            for page_num, image in enumerate(images, 1):
                temp_image_path = f"/tmp/page_{page_num}_{uuid.uuid4().hex}.png"
                image.save(temp_image_path)
                self.temp_files.append(temp_image_path)
                
                ocr_result = self.extract_text_with_ocr(temp_image_path)
                
                if ocr_result['text'].strip():
                    ws = wb.create_sheet(title=f"Page_{page_num}")
                    
                    lines = ocr_result.get('lines', [])
                    if lines:
                        for row_idx, line in enumerate(lines, 1):
                            text = line.get('text', '').strip()
                            if text:
                                cells = [cell.strip() for cell in text.split('\t')]
                                if len(cells) == 1:
                                    cells = [cell.strip() for cell in text.split('  ') if cell.strip()]
                                
                                for col_idx, cell_value in enumerate(cells, 1):
                                    ws.cell(row=row_idx, column=col_idx, value=cell_value)
                    else:
                        lines = ocr_result['text'].split('\n')
                        for row_idx, line in enumerate(lines, 1):
                            if line.strip():
                                ws.cell(row=row_idx, column=1, value=line.strip())
                    
                    total_confidence += ocr_result['confidence']
                    pages_processed += 1
            
            if pages_processed == 0:
                ws = wb.create_sheet(title="No_Content_Found")
                ws.cell(row=1, column=1, value="No readable content found in PDF")
            
            wb.save(output_path)
            
            avg_confidence = total_confidence / pages_processed if pages_processed > 0 else 0
            
            return {
                'success': True,
                'pages_processed': pages_processed,
                'quality_score': avg_confidence,
                'method': 'ocr_fallback'
            }
            
        except Exception as e:
            print(f"PDF to Excel OCR conversion failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def convert_excel_to_pdf(self, excel_path: str, output_path: str) -> Dict[str, Any]:
        """Convert Excel to PDF with professional formatting"""
        try:
            wb = load_workbook(excel_path)
            
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=12,
                textColor=colors.HexColor('#FF6B35'),
                alignment=1  # Center alignment
            )
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                
                story.append(Paragraph(f"Sheet: {sheet_name}", title_style))
                
                data = []
                for row in ws.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        data.append([str(cell) if cell is not None else '' for cell in row])
                
                if data:
                    table = Table(data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FF6B35')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            
            return {
                'success': True,
                'sheets_processed': len(wb.sheetnames),
                'quality_score': 95
            }
            
        except Exception as e:
            print(f"Excel to PDF conversion failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def convert_to_csv(self, file_path: str, output_path: str, file_type: str) -> Dict[str, Any]:
        """Convert various formats to CSV with structure preservation"""
        try:
            if file_type in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path, sheet_name=0)
                df.to_csv(output_path, index=False)
                return {'success': True, 'rows': len(df), 'columns': len(df.columns)}
                
            elif file_type == '.pdf':
                tables = self.extract_tables_from_pdf(file_path)
                if tables:
                    best_table = max(tables, key=lambda x: x.get('accuracy', 0))
                    df = pd.DataFrame(best_table['data'])
                    df.to_csv(output_path, index=False)
                    return {'success': True, 'rows': len(df), 'columns': len(df.columns)}
                else:
                    return {'success': False, 'error': 'No tables found in PDF'}
                    
            elif file_type in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif', '.webp']:
                ocr_result = self.extract_text_with_ocr(file_path)
                lines = ocr_result.get('lines', [])
                
                csv_data = []
                for line in lines:
                    text = line.get('text', '').strip()
                    if text:
                        for delimiter in [',', '\t', '  ', '|', ';']:
                            row = [cell.strip() for cell in text.split(delimiter) if cell.strip()]
                            if len(row) > 1:
                                csv_data.append(row)
                                break
                
                if csv_data:
                    max_cols = max(len(row) for row in csv_data)
                    normalized_data = []
                    for row in csv_data:
                        normalized_row = row + [''] * (max_cols - len(row))
                        normalized_data.append(normalized_row)
                    
                    headers = normalized_data[0] if normalized_data else []
                    data_rows = normalized_data[1:] if len(normalized_data) > 1 else normalized_data
                    
                    df = pd.DataFrame(data_rows, columns=headers if headers else None)
                    df.to_csv(output_path, index=False)
                    return {'success': True, 'rows': len(df), 'columns': len(df.columns)}
                else:
                    return {'success': False, 'error': 'No tabular data detected in image'}
            
            elif file_type == '.docx':
                doc = DocxDocument(file_path)
                text_data = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_data.append([paragraph.text.strip()])
                
                if text_data:
                    df = pd.DataFrame(text_data, columns=['Content'])
                    df.to_csv(output_path, index=False)
                    return {'success': True, 'rows': len(df), 'columns': 1}
                else:
                    return {'success': False, 'error': 'No content found in document'}
            
            return {'success': False, 'error': f'Unsupported conversion: {file_type} to CSV'}
            
        except Exception as e:
            print(f"CSV conversion failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def convert_to_word(self, file_path: str, output_path: str, file_type: str) -> Dict[str, Any]:
        """Convert various formats to Word document"""
        try:
            doc = DocxDocument()
            
            if file_type == '.pdf':
                images = convert_from_path(file_path, dpi=200)
                total_confidence = 0
                pages_processed = 0
                
                for page_num, image in enumerate(images, 1):
                    temp_image_path = f"/tmp/pdf_page_{page_num}_{uuid.uuid4().hex}.png"
                    image.save(temp_image_path)
                    self.temp_files.append(temp_image_path)
                    
                    ocr_result = self.extract_text_with_ocr(temp_image_path)
                    
                    if ocr_result['text'].strip():
                        heading = doc.add_heading(f'Page {page_num}', level=1)
                        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        lines = ocr_result.get('lines', [])
                        for line in lines:
                            text = line.get('text', '').strip()
                            if text:
                                paragraph = doc.add_paragraph(text)
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
                        total_confidence += ocr_result['confidence']
                        pages_processed += 1
                
                avg_confidence = total_confidence / pages_processed if pages_processed > 0 else 0
                
            elif file_type in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path, sheet_name=None)
                
                for sheet_name, sheet_df in df.items():
                    heading = doc.add_heading(f'Sheet: {sheet_name}', level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if not sheet_df.empty:
                        table = doc.add_table(rows=1, cols=len(sheet_df.columns))
                        table.style = 'Table Grid'
                        
                        hdr_cells = table.rows[0].cells
                        for i, column in enumerate(sheet_df.columns):
                            hdr_cells[i].text = str(column)
                            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                        
                        for _, row in sheet_df.iterrows():
                            row_cells = table.add_row().cells
                            for i, value in enumerate(row):
                                row_cells[i].text = str(value) if pd.notna(value) else ''
                
                avg_confidence = 95
                
            elif file_type == '.csv':
                df = pd.read_csv(file_path)
                
                heading = doc.add_heading('CSV Data', level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if not df.empty:
                    table = doc.add_table(rows=1, cols=len(df.columns))
                    table.style = 'Table Grid'
                    
                    hdr_cells = table.rows[0].cells
                    for i, column in enumerate(df.columns):
                        hdr_cells[i].text = str(column)
                        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    
                    for _, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for i, value in enumerate(row):
                            row_cells[i].text = str(value) if pd.notna(value) else ''
                
                avg_confidence = 95
                
            else:
                return {'success': False, 'error': f'Unsupported conversion: {file_type} to Word'}
            
            doc.save(output_path)
            
            return {
                'success': True,
                'quality_score': avg_confidence,
                'method': 'comprehensive_word_conversion'
            }
            
        except Exception as e:
            print(f"Word conversion failed: {e}")
            return {'success': False, 'error': str(e)}

processor = EnterpriseDocumentProcessor()

@router.get("/docusinage/formats")
async def get_supported_formats():
    """Get comprehensive list of supported input and output formats"""
    return {
        'input_formats': SUPPORTED_FORMATS['input'],
        'output_formats': SUPPORTED_FORMATS['output'],
        'ocr_available': OCR_AVAILABLE,
        'max_file_size_mb': MAX_FILE_SIZE // (1024 * 1024),
        'features': [
            'Enterprise-grade OCR',
            'Table structure preservation',
            'Multi-format conversion',
            'Bidirectional conversion',
            'Advanced image enhancement',
            'Multiple extraction methods'
        ]
    }

@router.post("/docusinage/upload")
async def upload_document(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload document for enterprise processing"""
    try:
        if file.size and file.size > MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail=f"File too large. Maximum size: {MAX_FILE_SIZE // (1024 * 1024)}MB")
        
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in SUPPORTED_FORMATS['input']:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Supported: {', '.join(SUPPORTED_FORMATS['input'])}"
            )
        
        file_id = str(uuid.uuid4())
        filename = f"{file_id}_{file.filename}"
        file_path = os.path.join(UPLOAD_DIR, filename)
        
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        mime_type, _ = mimetypes.guess_type(file_path)
        
        document_info = {
            'id': file_id,
            'filename': file.filename,
            'file_path': file_path,
            'file_type': file_ext,
            'mime_type': mime_type,
            'file_size': os.path.getsize(file_path),
            'uploaded_at': datetime.utcnow().isoformat(),
            'user_id': current_user['id'],
            'status': 'uploaded',
            'processing_status': 'pending'
        }
        
        return {
            'success': True,
            'document': document_info,
            'message': 'Document uploaded successfully'
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.post("/docusinage/{document_id}/process")
async def process_document(
    document_id: str,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Process document with enterprise OCR and analysis"""
    try:
        document_files = [f for f in os.listdir(UPLOAD_DIR) if f.startswith(document_id)]
        if not document_files:
            raise HTTPException(status_code=404, detail="Document not found")
        
        file_path = os.path.join(UPLOAD_DIR, document_files[0])
        file_ext = os.path.splitext(document_files[0])[1].lower()
        
        processing_result = {
            'document_id': document_id,
            'processing_started': datetime.utcnow().isoformat(),
            'file_type': file_ext,
            'status': 'processing'
        }
        
        if file_ext == '.pdf':
            try:
                tables = processor.extract_tables_from_pdf(file_path)
                
                images = convert_from_path(file_path, dpi=200)
                ocr_results = []
                total_confidence = 0
                
                for page_num, image in enumerate(images, 1):
                    temp_image_path = f"/tmp/pdf_page_{page_num}_{uuid.uuid4().hex}.png"
                    image.save(temp_image_path)
                    processor.temp_files.append(temp_image_path)
                    
                    ocr_result = processor.extract_text_with_ocr(temp_image_path)
                    ocr_results.append({
                        'page': page_num,
                        'text': ocr_result['text'],
                        'confidence': ocr_result['confidence'],
                        'word_count': ocr_result['total_words']
                    })
                    total_confidence += ocr_result['confidence']
                
                avg_confidence = total_confidence / len(images) if images else 0
                
                processing_result.update({
                    'status': 'completed',
                    'pages_processed': len(images),
                    'tables_found': len(tables),
                    'ocr_results': ocr_results,
                    'overall_confidence': avg_confidence,
                    'tables': tables[:5],  # Limit response size
                    'processing_completed': datetime.utcnow().isoformat()
                })
                
            except Exception as e:
                processing_result.update({
                    'status': 'failed',
                    'error': str(e),
                    'processing_completed': datetime.utcnow().isoformat()
                })
                
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif', '.webp']:
            try:
                ocr_result = processor.extract_text_with_ocr(file_path)
                
                processing_result.update({
                    'status': 'completed',
                    'text': ocr_result['text'],
                    'confidence': ocr_result['confidence'],
                    'word_count': ocr_result['total_words'],
                    'words': ocr_result['words'][:100],  # Limit response size
                    'processing_completed': datetime.utcnow().isoformat()
                })
                
            except Exception as e:
                processing_result.update({
                    'status': 'failed',
                    'error': str(e),
                    'processing_completed': datetime.utcnow().isoformat()
                })
                
        elif file_ext in ['.xlsx', '.xls']:
            try:
                df = pd.read_excel(file_path, sheet_name=None)
                sheets_info = []
                
                for sheet_name, sheet_df in df.items():
                    sheets_info.append({
                        'name': sheet_name,
                        'rows': len(sheet_df),
                        'columns': len(sheet_df.columns),
                        'headers': sheet_df.columns.tolist()
                    })
                
                processing_result.update({
                    'status': 'completed',
                    'sheets': sheets_info,
                    'total_sheets': len(df),
                    'processing_completed': datetime.utcnow().isoformat()
                })
                
            except Exception as e:
                processing_result.update({
                    'status': 'failed',
                    'error': str(e),
                    'processing_completed': datetime.utcnow().isoformat()
                })
                
        elif file_ext == '.csv':
            try:
                df = pd.read_csv(file_path)
                
                processing_result.update({
                    'status': 'completed',
                    'rows': len(df),
                    'columns': len(df.columns),
                    'headers': df.columns.tolist(),
                    'sample_data': df.head().to_dict('records'),
                    'processing_completed': datetime.utcnow().isoformat()
                })
                
            except Exception as e:
                processing_result.update({
                    'status': 'failed',
                    'error': str(e),
                    'processing_completed': datetime.utcnow().isoformat()
                })
                
        elif file_ext == '.docx':
            try:
                doc = DocxDocument(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                
                processing_result.update({
                    'status': 'completed',
                    'paragraphs': len(paragraphs),
                    'content_preview': paragraphs[:10],  # First 10 paragraphs
                    'processing_completed': datetime.utcnow().isoformat()
                })
                
            except Exception as e:
                processing_result.update({
                    'status': 'failed',
                    'error': str(e),
                    'processing_completed': datetime.utcnow().isoformat()
                })
        
        else:
            processing_result.update({
                'status': 'completed',
                'message': 'File uploaded successfully, ready for conversion',
                'processing_completed': datetime.utcnow().isoformat()
            })
        
        processor.cleanup_temp_files()
        
        return processing_result
        
    except Exception as e:
        processor.cleanup_temp_files()
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

@router.post("/docusinage/{document_id}/convert")
async def convert_document(
    document_id: str,
    target_format: str = Query(..., description="Target format (pdf, xlsx, csv, docx, etc.)"),
    current_user: dict = Depends(get_current_user)
):
    """Convert document to specified format with structure preservation"""
    try:
        document_files = [f for f in os.listdir(UPLOAD_DIR) if f.startswith(document_id)]
        if not document_files:
            raise HTTPException(status_code=404, detail="Document not found")
        
        source_file = os.path.join(UPLOAD_DIR, document_files[0])
        source_ext = os.path.splitext(document_files[0])[1].lower()
        
        if not target_format.startswith('.'):
            target_format = f'.{target_format}'
        
        if target_format not in SUPPORTED_FORMATS['output']:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported target format. Supported: {', '.join(SUPPORTED_FORMATS['output'])}"
            )
        
        output_filename = f"{document_id}_converted{target_format}"
        output_path = os.path.join(PROCESSED_DIR, output_filename)
        
        conversion_result = {
            'document_id': document_id,
            'source_format': source_ext,
            'target_format': target_format,
            'conversion_started': datetime.utcnow().isoformat()
        }
        
        if target_format == '.xlsx':
            if source_ext == '.pdf':
                result = processor.convert_pdf_to_excel(source_file, output_path)
            elif source_ext == '.csv':
                df = pd.read_csv(source_file)
                df.to_excel(output_path, index=False)
                result = {'success': True, 'quality_score': 95}
            else:
                result = {'success': False, 'error': f'Conversion from {source_ext} to {target_format} not supported'}
                
        elif target_format == '.pdf':
            if source_ext in ['.xlsx', '.xls']:
                result = processor.convert_excel_to_pdf(source_file, output_path)
            elif source_ext == '.csv':
                df = pd.read_csv(source_file)
                wb = Workbook()
                ws = wb.active
                for r_idx, row in enumerate(df.itertuples(index=False), 1):
                    for c_idx, value in enumerate(row, 1):
                        ws.cell(row=r_idx, column=c_idx, value=value)
                temp_excel = f"/tmp/{document_id}_temp.xlsx"
                wb.save(temp_excel)
                result = processor.convert_excel_to_pdf(temp_excel, output_path)
                os.remove(temp_excel)
            else:
                result = {'success': False, 'error': f'Conversion from {source_ext} to {target_format} not supported'}
                
        elif target_format == '.csv':
            result = processor.convert_to_csv(source_file, output_path, source_ext)
            
        elif target_format == '.docx':
            result = processor.convert_to_word(source_file, output_path, source_ext)
            
        else:
            result = {'success': False, 'error': f'Target format {target_format} not implemented yet'}
        
        if result.get('success'):
            conversion_result.update({
                'status': 'completed',
                'output_file': output_filename,
                'file_size': os.path.getsize(output_path),
                'quality_score': result.get('quality_score', 0),
                'conversion_completed': datetime.utcnow().isoformat()
            })
        else:
            conversion_result.update({
                'status': 'failed',
                'error': result.get('error', 'Unknown conversion error'),
                'conversion_completed': datetime.utcnow().isoformat()
            })
        
        processor.cleanup_temp_files()
        return conversion_result
        
    except Exception as e:
        processor.cleanup_temp_files()
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

@router.get("/docusinage/{document_id}/download/{filename}")
async def download_converted_file(
    document_id: str,
    filename: str,
    current_user: dict = Depends(get_current_user)
):
    """Download converted file"""
    try:
        file_path = os.path.join(PROCESSED_DIR, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")
        
        if not filename.startswith(document_id):
            raise HTTPException(status_code=403, detail="Access denied")
        
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type='application/octet-stream'
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

@router.delete("/docusinage/{document_id}")
async def delete_document(
    document_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete document and all associated files"""
    try:
        deleted_files = []
        
        for filename in os.listdir(UPLOAD_DIR):
            if filename.startswith(document_id):
                file_path = os.path.join(UPLOAD_DIR, filename)
                os.remove(file_path)
                deleted_files.append(filename)
        
        for filename in os.listdir(PROCESSED_DIR):
            if filename.startswith(document_id):
                file_path = os.path.join(PROCESSED_DIR, filename)
                os.remove(file_path)
                deleted_files.append(filename)
        
        if not deleted_files:
            raise HTTPException(status_code=404, detail="Document not found")
        
        return {
            'success': True,
            'document_id': document_id,
            'deleted_files': deleted_files,
            'deleted_at': datetime.utcnow().isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Deletion failed: {str(e)}")

@router.get("/docusinage/health")
async def health_check():
    """Health check for Docusinage service"""
    return {
        'status': 'healthy',
        'ocr_available': OCR_AVAILABLE,
        'upload_dir_exists': os.path.exists(UPLOAD_DIR),
        'processed_dir_exists': os.path.exists(PROCESSED_DIR),
        'timestamp': datetime.utcnow().isoformat()
    }
